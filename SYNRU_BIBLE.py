#!/usr/bin/env python3
# coding=utf-8

import os
import re
import sqlite3
import subprocess

import psutil
import win32com.client

import time

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Номер версии издания
version = "(SYNRU)"
debug = False  # Режим отладки


# ----------------------------------------------------------------------------------------------------
# Функция для извлечения данных из базы данных SQLite в список кортежей (каждая строка - кортеж)
def extract_data_from_database(database_file):
    conn = sqlite3.connect(database_file)  # Подключение к базе данных
    cursor = conn.cursor()  # Создание курсора
    cursor.execute(
        "SELECT bn,short_name,book,chapter,verse,bible_text"
        " FROM COMPLETE_BIBLE ORDER BY bn,chapter,verse")  # Выполнение запроса
    data = cursor.fetchall()  # Получение всех строк результата запроса
    conn.close()  # Закрытие соединения
    return data  # Возврат результата


# ----------------------------------------------------------------------------------------------------
# Функция для создания .docx файла с тегами из данных Библии
def create_docx_file_with_tags(data, output_file):
    doc = Document()  # Создание документа Word

    style = doc.styles.add_style('SingleSpacing12', WD_STYLE_TYPE.PARAGRAPH)  # Создание стиля
    style.paragraph_format.line_spacing = 1.0  # Установка межстрочного интервала в 1.0
    style.paragraph_format.space_after = Pt(0)  # Установка отступа снизу в 0 пунктов
    style.paragraph_format.space_before = Pt(0)  # Установка отступа сверху в 0 пунктов
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Выравнивание слева
    style.font.size = Pt(12)  # Размер шрифта 12 пунктов

    style = doc.styles.add_style('SingleSpacing11', WD_STYLE_TYPE.PARAGRAPH)  # Создание стиля
    style.paragraph_format.line_spacing = 1.0  # Установка межстрочного интервала в 1.0
    style.paragraph_format.space_after = Pt(0)  # Установка отступа снизу в 0 пунктов
    style.paragraph_format.space_before = Pt(0)  # Установка отступа сверху в 0 пунктов
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Выравнивание слева
    style.font.size = Pt(11)  # Размер шрифта 11 пунктов

    style = doc.styles.add_style('SingleSpacing10', WD_STYLE_TYPE.PARAGRAPH)  # Создание стиля
    style.paragraph_format.line_spacing = 1.0  # Установка межстрочного интервала в 1.0
    style.paragraph_format.space_after = Pt(0)  # Установка отступа снизу в 0 пунктов
    style.paragraph_format.space_before = Pt(0)  # Установка отступа сверху в 0 пунктов
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Выравнивание слева
    style.font.size = Pt(10)  # Размер шрифта 10 пунктов

    current_book = None  # Текущая книга
    current_book_short_name = None  # Текущее сокращенное наименование книги

    current_chapter = None  # Текущая глава

    for row in data:  # Обработка каждой строки данных
        book = row[2]  # Название книги Библии
        book_short_name = row[1]  # Сокращенное наименование книги

        chapter = row[3]  # Номер главы

        if debug:
            if ((book != "Бытие") and (book != "Евангелие от Марка")):
                continue

        # Если книга изменилась
        if book != current_book:
            print(f"   ⌛️ Обработка книги № {row[0]}: {book} ...")

            paragraph = doc.add_paragraph(style='SingleSpacing12')  # Добавление параграфа
            paragraph_format = paragraph.paragraph_format  # Получение форматирования параграфа
            paragraph_format.space_after = Pt(7)  # Установка отступа снизу в 7 пунктов
            paragraph_format.space_before = Pt(12)  # Установка отступа сверху в 12 пунктов
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Выравнивание по центру
            paragraph.style = doc.styles['Heading 1']  # Установка стиля "Заголовок 1"

            # Добавление названия книги
            paragraph.add_run(f"{book}")  # ************

        # Если глава изменилась или книга изменилась
        if chapter != current_chapter or book != current_book:
            paragraph = doc.add_paragraph(style='SingleSpacing12')  # Добавление параграфа
            paragraph_format = paragraph.paragraph_format  # Получение форматирования параграфа
            paragraph_format.space_after = Pt(7)  # Установка отступа снизу в 7 пунктов
            paragraph_format.space_before = Pt(12)  # Установка отступа сверху в 12 пунктов
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Выравнивание по центру
            paragraph.style = doc.styles['Heading 2']  # Установка стиля "Заголовок 2"

            # Добавление номера главы
            paragraph.add_run(f"Глава {chapter}")  # ************

            current_chapter = chapter  # Обновление номера текущей главы

        if book != current_book:  # Если книга изменилась
            current_book = book  # Обновление наименования текущей книги
            current_book_short_name = book_short_name  # Обновление сокращенного наименования текущей книги

        # Добавление номера стиха и текста стиха
        cleaned_row = re.sub(" +", " ", row[5])  # Очистка текста стиха от лишних пробелов
        paragraph = doc.add_paragraph(style='SingleSpacing12')  # Добавление параграфа
        paragraph_format = paragraph.paragraph_format  # Получение форматирования параграфа
        paragraph_format.space_before = Pt(0)  # Установка отступа сверху в 0 пунктов
        paragraph_format.space_after = Pt(0)  # Установка отступа снизу в 0 пунктов
        paragraph_format.line_spacing = 1.0  # Установка межстрочного интервала в 1.0
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Выравнивание слева
        paragraph.style = doc.styles['Normal']  # Установка стиля "Обычный"
        paragraph.style.font.size = Pt(12)  # Размер шрифта 12 пунктов

        # Добавление номера стиха и текста стиха в параграф
        paragraph.add_run( f"<b><sup><small>{row[4]}</small></sup></b> {cleaned_row}")

    doc.save(output_file)  # Сохранение документа


# ----------------------------------------------------------------------------------------------------
# Функция для конвертации .docx файла в .html
def convert_docx_to_html(input_file, output_file):
    word = win32com.client.Dispatch('Word.Application')
    doc = word.Documents.Open(input_file)
    doc.SaveAs(output_file, FileFormat=10)  # 10 - это код для HTML
    doc.Close()
    word.Quit()
    time.sleep(1)
    for proc in psutil.process_iter(attrs=['pid', 'name']):
        if proc.info['name'].lower() == 'winword.exe':
            return False
    return True


# ----------------------------------------------------------------------------------------------------
# Функция для замены тегов в файле .html на символы < и >  и сохранения изменений в том же файле
def replace_tags_in_html_file(file_path):
    # Открытие файла для чтения
    with open(file_path, 'r', encoding='windows-1251') as file:
        file_content = file.read()

    # Замена тегов
    file_content = file_content.replace('&lt;', '<')  # Замена символа <
    file_content = file_content.replace('&gt;', '>')  # Замена символа >
    file_content = file_content.replace('&quot;', '"')  # Двойные кавычки
    file_content = file_content.replace('&#8202;', ' ')  # Тонкий пробел

    file_content = re.sub(' +', ' ', file_content)  # Удаление лишних пробелов

    # Открытие файла для записи
    with open(file_path, 'w', encoding='windows-1251') as file:
        file.write(file_content)


# ----------------------------------------------------------------------------------------------------
# Функция для конвертации .html файла в .docx
def convert_html_to_docx(input_file, output_file):
    word = win32com.client.Dispatch('Word.Application')
    doc = word.Documents.Open(input_file)
    doc.SaveAs(output_file, FileFormat=16)  # 16 - это код для DOCX
    doc.Close()
    word.Quit()
    time.sleep(1)
    for proc in psutil.process_iter(attrs=['pid', 'name']):
        if proc.info['name'].lower() == 'winword.exe':
            return False
    return True


# ----------------------------------------------------------------------------------------------------
# Главная функция программы
def main():
    database_file = "SYNRU_BIBLE.sqlite"  # Имя файла базы данных SQLite с данными Библии
    docx_file_tags = os.path.abspath(
        f"docx_file_tags{version}.docx")  # Имя промежуточного док. Word без тегов
    html_file = os.path.abspath(
        f"БИБЛИЯ {version}.html")  # Имя готового файла Библии формата html
    docx_file = os.path.abspath(
        f"БИБЛИЯ {version}.docx")  # Имя готового файла Библии формата odt

    if debug:
        print("(!) Включен режим отладки !\n")

    print(f"Начато создание Библии {version}:")

    print(f"[1/5] Извлечение данных из БД ...")
    database_data = extract_data_from_database(database_file)  # Извлечение данных из базы данных

    print(f"[2/5] Создание файла с данными ...")
    create_docx_file_with_tags(database_data, docx_file_tags)  # Создание Библии из данных и сохранение в файле

    print(f"[3/5] Конвертирование ...")
    convert_docx_to_html(docx_file_tags, html_file)  # Конвертация файла .docx в .html

    print(f"[4/5] Замена тегов ...")
    replace_tags_in_html_file(html_file)  # Замена тегов в файле .html

    print(f"[5/5] Финальное конвертирование ...")
    convert_html_to_docx(html_file, docx_file)  # Конвертация файла .html в .docx
    os.remove(docx_file_tags)  # Удаление промежуточного файла docx без тегов

    print(f"Создание Библии {version} успешно завершено !\n")

    subprocess.Popen([docx_file], shell=True)  # Открытие готового файла Библии формата docx в Word


# ----------------------------------------------------------------------------------------------------


if __name__ == "__main__":
    main()
